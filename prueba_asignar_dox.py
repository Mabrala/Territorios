from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# --- Función para escribir en la celda con Arial 11 y centrado ---
def set_cell_text(cell, text):
    """Escribe texto en una celda con fuente Arial, tamaño 11 y centrado."""
    cell.text = ""  # limpiar contenido previo
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # centrar texto
    run = paragraph.add_run(text)
    
    # Fuente y tamaño fijos
    run.font.name = "Arial"
    run.font.size = Pt(11)
    
    # Garantizar que Word reconozca Arial
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), "Arial")

# --- Datos y ruta del documento ---
document_path = os.path.join(os.path.dirname(__file__), 'AA 25-26.docx')
document = Document(document_path)

territory_code = "O2"
assigned_to = "Juan Perez"
assigned_date = "15/10/24"
updated = False

# --- Recorrer tablas para asignar nombre y fecha ---
for table in document.tables:
    for row_index, row in enumerate(table.rows):
        first_cell = row.cells[0].text.strip()
        
        if first_cell == territory_code:
            # Buscar primera celda vacía en "Asignado a"
            for i in range(2, len(row.cells)):
                if not row.cells[i].text.strip():
                    # Escribir nombre en la celda correspondiente
                    set_cell_text(row.cells[i], assigned_to)
                    
                    # Escribir fecha en la fila siguiente, misma columna
                    if row_index + 1 < len(table.rows):
                        next_row = table.rows[row_index + 1]
                        set_cell_text(next_row.cells[i], assigned_date)
                    
                    updated = True
                    break
        if updated:
            break
    if updated:
        break

# --- Guardar cambios ---
document.save(document_path)
