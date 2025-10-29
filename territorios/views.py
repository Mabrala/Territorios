from django.contrib import messages
from django.shortcuts import render, redirect
import os
os.environ['OAUTHLIB_INSECURE_TRANSPORT'] = '1'
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from .models import *
from google_auth_oauthlib.flow import Flow
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger

#para imagenes
from django.http import HttpResponse, Http404
from googleapiclient.http import MediaIoBaseDownload, MediaIoBaseUpload
from googleapiclient.errors import HttpError
import io
import mimetypes

#escribit docx y subirlo
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
                
#Revisa las credenciales, si no las tiene devuelve None, si las tine devuelve lista
def check_creds(request):
    return request.session.get('credentials')
    
def clasify_items(items):
    for item in items:
        if item["mimeType"] == "application/vnd.google-apps.folder":
            item["type"] = "folder"
        else:
            # Clasificar archivos por tipo
            if item["mimeType"] in [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword"
            ]:
                item["type"] = "docx"
            elif item["mimeType"] == "application/pdf":
                item["type"] = "pdf"
            elif item["mimeType"] in [
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                "application/vnd.ms-excel"
            ]:
                item["type"] = "excel"
            elif item["mimeType"].startswith("image/"):
                item["type"] = "image"
            else:
                item["type"] = "other"
    return items

def paginate_items(request, items, per_page=11):
    # Ordena primero carpetas, luego archivos por nombre
    items_sorted = sorted(
        items,
        key=lambda x: (x.get("type") != "folder", x.get("name", "").lower())
    )
    page = request.GET.get('page', 1)
    paginator = Paginator(items_sorted, per_page)
    try:
        files = paginator.page(page)
    except PageNotAnInteger:
        files = paginator.page(1)
    except EmptyPage:
        files = paginator.page(paginator.num_pages)
    return files

def index(request):
    creds_info = check_creds(request)
    if creds_info:
        drive_folder = Folder.objects.first()
        if drive_folder != None:
            creds = Credentials(**creds_info)
            service = build('drive', 'v3', credentials=creds)

            results = service.files().list(
                q=f"'{drive_folder.id_folder}' in parents",
                pageSize=100,
                fields="files(id, name, mimeType)",
                orderBy="folder,name",
             ).execute()
            
            items = results.get("files", [])
            items = clasify_items(items)

            files = paginate_items(request, items)
            return render(request, 'index.html', {"files": files})
    else:
        return redirect("drive_auth_init")


#Lista el contenido de cualquier carpeta
def list_folder_content(request,folder_id):
    creds_info = check_creds(request)
    if creds_info:
        creds = Credentials(**creds_info)
        service = build('drive', 'v3', credentials=creds)
        folder = service.files().get(
            fileId=folder_id,
            fields="name"
        ).execute()
        results = service.files().list(
                q=f"'{folder_id}' in parents",
                fields="files(id, name, mimeType)",
                orderBy="folder,name",
             ).execute()
        
        items = results.get("files", [])
        items = clasify_items(items)

        files = paginate_items(request, items)
        return render(request, 'folder/list.html', {"files": files, "folder_name": folder})   
    else:
        return redirect("drive_auth_init")

#
# --- Función para escribir en la celda con Arial 11 y centrado ---
def set_cell_text(cell, text):
    """Escribe texto en una celda con fuente Arial, tamaño 11 y centrado."""
    cell.text = ""  # limpiar contenido previo
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # centrar texto
    run = paragraph.add_run(text)
    
    # Fuente y tamaño fijos
    run.font.name = "Arial"
    run.font.size = Pt(11)
    
    # Garantizar que Word reconozca Arial
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), "Arial")

#Funcion para asignar en excel
from openpyxl import load_workbook
def actualizar_excel_drive(service, excel_id, codigo, nombre, fecha):

    # --- Descargar Excel desde Drive ---
    excel_request = service.files().get_media(fileId=excel_id)
    excel_bytes = io.BytesIO()
    downloader = MediaIoBaseDownload(excel_bytes, excel_request)
    done = False
    while not done:
        status, done = downloader.next_chunk()
    excel_bytes.seek(0)

    # --- Modificar con openpyxl ---
    wb = load_workbook(excel_bytes)

    # --- Borrar de RECIBIDOS ---
    if "RECIBIDOS" in wb.sheetnames:
        hoja_rec = wb["RECIBIDOS"]
        fila_a_borrar = None
        for fila in range(1, hoja_rec.max_row + 1):
            celda = hoja_rec.cell(row=fila, column=1).value
            if celda == codigo:
                fila_a_borrar = fila
                break
        if fila_a_borrar:
            hoja_rec.delete_rows(fila_a_borrar)

    # --- Añadir a ENTREGADOS ---
    hoja = wb["ENTREGADOS"]

    fila = hoja.max_row + 1
    while all(cell.value is None for cell in hoja[fila]):
        fila -= 1
    
    hoja.cell(row=fila, column=1, value=codigo)
    hoja.cell(row=fila, column=2, value=nombre)
    hoja.cell(row=fila, column=3, value=fecha)

    # --- Guardar cambios en memoria ---
    updated_excel = io.BytesIO()
    wb.save(updated_excel)
    updated_excel.seek(0)
    wb.close()

    # --- 3️⃣ Subir de nuevo al mismo archivo en Drive ---
    media = MediaIoBaseUpload(
        updated_excel,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        resumable=True
    )
    service.files().update(fileId=excel_id, media_body=media).execute()


#Metodo para asignar territorios en docx
import re
def assign_territory(request, file_name):
    creds_info = check_creds(request)
    if not creds_info:
        return redirect("drive_auth_init")

    creds = Credentials(**creds_info)
    service = build("drive", "v3", credentials=creds)
    
    if request.method == "POST":
        
        assigned_to = request.POST.get("assigned_to")
        assigned_date = request.POST.get("assigned_date")
        regex_date = re.compile(r"^\d{2}\/\d{2}\/\d{2}$")
        
        if not assigned_to:
            message = "El campo 'Asignar a' es obligatorio."
            messages.error(request, message)
            referer = request.META.get('HTTP_REFERER') or '/'
            return redirect(referer)
        
        if not regex_date.match(assigned_date):
            message = "El campo 'Fecha' debe tener el formato DD-MM-AA."
            messages.error(request, message)
            referer = request.META.get('HTTP_REFERER') or '/'
            return redirect(referer)
        
        drive_folder = Folder.objects.first()
        
        register_id = drive_folder.register_id
        
        # Convertir a código de territorio (R-4.png/jpg/pdf → R4)
        #territory_code = re.sub(r"[-_.].*", "", file_name.split('.')[0]) 
        territory_code = file_name.replace(".png","").replace(".jpg","").replace(".pdf","").replace("-","").replace(".","")
        
        
        register_request = service.files().get_media(fileId=register_id)
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, register_request)
        done = False
        while not done:
            status, done = downloader.next_chunk()
        fh.seek(0)
        
        # Modificar el DOCX con python-docx
        document = Document(fh)
        updated = False
        
# --- Recorrer tablas para asignar nombre y fecha ---
        for table in document.tables:
            for row_index, row in enumerate(table.rows):
                first_cell = row.cells[0].text.strip()
                
                if first_cell == territory_code:
                    # Buscar primera celda vacía en "Asignado a"
                    for i in range(2, len(row.cells)):
                        if not row.cells[i].text.strip():
                            # Escribir nombre en la celda correspondiente
                            set_cell_text(row.cells[i], assigned_to)
                            
                            # Escribir fecha en la fila siguiente, misma columna
                            if row_index + 1 < len(table.rows):
                                next_row = table.rows[row_index + 1]
                                set_cell_text(next_row.cells[i], assigned_date)
                            
                            updated = True
                            break
                if updated:
                    break
            if updated:
                break

        if not updated:
            message = f"Ha ocurrido un error al asignar el territorio {territory_code}."
            messages.error(request, message)
            referer = request.META.get('HTTP_REFERER') or '/'
            return redirect(referer)

        # Guardar el documento modificado en memoria
        updated_fh = io.BytesIO()
        document.save(updated_fh)
        updated_fh.seek(0)

        # Subirlo de nuevo a Drive (sobrescribiendo)
        media = MediaIoBaseUpload(updated_fh, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document", resumable=True)
        service.files().update(fileId=register_id, media_body=media).execute()
        
        try:
            excel_id = drive_folder.ex_id
            actualizar_excel_drive(service, excel_id, territory_code, assigned_to, assigned_date)
        except Exception as e:
            print(f"⚠️ Error al actualizar Excel en Drive: {e}")

        message = f"{territory_code} asignado a {assigned_to} con fecha {assigned_date}."
        messages.success(request, message)
        referer = request.META.get('HTTP_REFERER') or '/'
        
        return redirect(referer)
    
    else:
        referer = request.META.get('HTTP_REFERER')
        if referer:
            return redirect(referer)
        return redirect('index')

def read_excel(request):
    creds_info = check_creds(request)
    if not creds_info:
        return redirect("drive_auth_init")

    drive_folder = Folder.objects.first()
    excel_id = drive_folder.ex_id
    creds = Credentials(**creds_info)
    service = build('drive', 'v3', credentials=creds)

    # Descargar el archivo Excel
    excel_request = service.files().get_media(fileId=excel_id)
    excel_bytes = io.BytesIO()
    downloader = MediaIoBaseDownload(excel_bytes, excel_request)
    done = False
    while not done:
        status, done = downloader.next_chunk()
    excel_bytes.seek(0)

    # Leer el contenido con openpyxl
    wb = load_workbook(excel_bytes)
    hoja = wb["ENTREGADOS"]
    data = []
    for fila in range(1, hoja.max_row + 1):
        codigo = hoja.cell(row=fila, column=1).value
        nombre = hoja.cell(row=fila, column=2).value
        fecha = hoja.cell(row=fila, column=3).value
        if codigo or nombre or fecha:
            data.append({
                "codigo": codigo,
                "nombre": nombre,
                "fecha": fecha
            })
    wb.close()

    # Insertar/actualizar registros en la base de datos para cada fila leida
    try:
        entregados_list = Entregados.objects.all()
        entregados_list.delete()
        for row in data:
            codigo = row.get("codigo") or ""
            nombre = row.get("nombre") or ""
            fecha = row.get("fecha") or ""
            if codigo != "":
                # Actualiza si ya existe un registro con ese código, o crea uno nuevo
                Entregados.objects.update_or_create(
                    territory=str(codigo),
                    defaults={
                        "brother": str(nombre),
                        "date": str(fecha)
                    }
                )
    except Exception:
        # No interrumpir la vista en caso de error con la BD
        pass

def entregados(request):
    creds_info = check_creds(request)
    if not creds_info:
        return redirect("drive_auth_init")
    read_excel(request)
    entregados_list = Entregados.objects.all().order_by('territory')
    return render(request, "entregados/entregados.html", {"entregados": entregados_list})

from datetime import datetime
def recibir(request):
    creds_info = check_creds(request)
    if not creds_info:
        return redirect("drive_auth_init")
    
    drive_folder = Folder.objects.first()
    excel_id = drive_folder.ex_id
    docx_id = drive_folder.register_id
    creds = Credentials(**creds_info)
    service = build('drive', 'v3', credentials=creds)
    
    if request.method == "POST":
        regex_date = re.compile(r"^\d{2}\/\d{2}\/\d{2}$")
        territory = request.POST.get("territory")
        date = datetime.strptime(request.POST['date'], "%Y-%m-%d").strftime("%d/%m/%y")
        assigned_to = request.POST.get("brother")
        
        if not regex_date.match(date):
            message = "El campo 'Fecha' debe tener el formato DD-MM-AA."
            messages.error(request, message)
            referer = request.META.get('HTTP_REFERER') or '/'
            return redirect(referer)
        
        register_request = service.files().get_media(fileId=docx_id)
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, register_request)
        done = False
        while not done:
            status, done = downloader.next_chunk()
        fh.seek(0)
        
        # Modificar el DOCX con python-docx
        document = Document(fh)
        updated = False
        
# --- Recorrer tablas para asignar nombre y fecha ---
        for table in document.tables:
            for row_index, row in enumerate(table.rows):
                first_cell = row.cells[0].text.strip()
                if first_cell == territory:
                    # Buscar primera celda vacía en "Asignado a"
                    for i in range(2, len(row.cells)):
                        if row.cells[i].text.strip() == assigned_to:
                            # Escribir fecha en la fila siguiente, columna+1
                            if row_index + 1 < len(table.rows):
                                next_row = table.rows[row_index + 1]
                                set_cell_text(next_row.cells[i+1], date)
                            
                            updated = True
                            break
                if updated:
                    break
            if updated:
                break
        
        excel_request = service.files().get_media(fileId=excel_id)
        excel_bytes = io.BytesIO()
        downloader = MediaIoBaseDownload(excel_bytes, excel_request)
        done_excel = False
        while not done_excel:
            status, done_excel = downloader.next_chunk()
        excel_bytes.seek(0)

        wb = load_workbook(excel_bytes)
        hoja = wb["ENTREGADOS"]
        excel_updated = False
        for fila in range(1, hoja.max_row + 1):
            if hoja.cell(row=fila, column=1).value == territory:
                hoja.cell(row=fila, column=1).value = ""
                hoja.cell(row=fila, column=2).value = ""
                hoja.cell(row=fila, column=3).value = ""
                excel_updated = True
                break
        
        if excel_updated:
            recibidos = wb["RECIBIDOS"]
            nueva_fila = recibidos.max_row + 1
            fila = hoja.max_row + 1
            while all(cell.value is None for cell in hoja[fila]):
                fila -= 1
            recibidos.cell(row=fila, column=1, value=territory)
            recibidos.cell(row=fila, column=2, value=date)
        wb.close()
        
        
        if updated and excel_updated:
            message = f"{territory} recibido de {assigned_to} con fecha {date}."
            messages.success(request, message)
            updated_fh = io.BytesIO()
            document.save(updated_fh)
            updated_fh.seek(0)
            updated_excel = io.BytesIO()
            wb.save(updated_excel)
            updated_excel.seek(0)
            
            media = MediaIoBaseUpload(updated_fh, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document", resumable=True)
            service.files().update(fileId=docx_id, media_body=media).execute()
            media = MediaIoBaseUpload(
                updated_excel,
                mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                resumable=True
            )
            service.files().update(fileId=excel_id, media_body=media).execute()
            bbdd = Entregados.objects.filter(territory=territory, brother=assigned_to).first()
            if bbdd:
                bbdd.delete()
        else:
            message = f"Ha ocurrido un error al recibir el territorio {territory}."
            messages.error(request, message)
        
        return redirect('entregados')


def search_in_folder(request, query):
    creds_info = check_creds(request)
    if not creds_info:
        return redirect("drive_auth_init")

    drive_folder = Folder.objects.first()
    folder_id = drive_folder.id_folder
    creds = Credentials(**creds_info)
    service = build('drive', 'v3', credentials=creds)

    folder = service.files().get(fileId=folder_id, fields="name").execute()

    def recursive_search(folder_id, query):
        results = []
        response = service.files().list(
            q=f"'{folder_id}' in parents and name contains '{query}'",
            pageSize=100,
            fields="files(id, name, mimeType)"
        ).execute()
        results.extend(response.get("files", []))
        subfolders = service.files().list(
            q=f"'{folder_id}' in parents and mimeType='application/vnd.google-apps.folder'",
            fields="files(id, name)"
        ).execute().get("files", [])
        for subfolder in subfolders:
            results.extend(recursive_search(subfolder["id"], query))
        return results

    items = recursive_search(folder_id, query)
    items = clasify_items(items)

    files = paginate_items(request, items)
    return render(request, 'folder/list.html', {"files": files, "folder_name": folder})


                
#Esta vista solo se utiliza para listar las carpetas de Drive cuando 
# se va a seleccionar la carpeta de trabajo
def list_drive_files(request):
    creds_info = check_creds(request)
    if creds_info:
        creds = Credentials(**creds_info)
        service = build('drive', 'v3', credentials=creds)

        # Mostrar solo carpetas de nivel superior (Mi unidad)
        results = service.files().list(
            q = "'root' in parents and mimeType='application/vnd.google-apps.folder'",
            pageSize=100,
            fields="files(id, name, mimeType)"
        ).execute()

        items = results.get('files', [])
        items = clasify_items(items)
        files = paginate_items(request, items)
        return render(request, "choose_drive_file/choose_drive_file.html", {"files": files})
    else:
        return redirect("drive_auth_init")


#visualizar imagenes
def view_file(request, file_id):
    creds_info = check_creds(request)
    if not creds_info:
        return redirect("drive_auth_init")

    creds = Credentials(**creds_info)
    service = build("drive", "v3", credentials=creds)

    try:
        # Obtener metadata del archivo
        file = service.files().get(
            fileId=file_id,
            fields="id, name, mimeType"
        ).execute()

        mime_type = file.get("mimeType", "")
        if not (mime_type.startswith("image/") or mime_type == "application/pdf"):
            raise Http404("El archivo no es una imagen ni un PDF")

        # Descargar el archivo
        request_file = service.files().get_media(fileId=file_id)
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, request_file)

        done = False
        while not done:
            status, done = downloader.next_chunk()

        fh.seek(0)  # reiniciar puntero

        # Determinar content type
        content_type = mime_type or mimetypes.guess_type(file["name"])[0] or "application/octet-stream"

        return HttpResponse(fh.read(), content_type=content_type)

    except Exception as e:
        raise Http404(f"No se pudo mostrar el archivo: {str(e)}")



#Selecciona la carpeta sobre la que se trabajara
def select_drive_folder(request, id_folder):
    creds_info = check_creds(request)
    if creds_info:
        creds = Credentials(**creds_info)
        service = build('drive', 'v3', credentials=creds)
        folder = service.files().get(
            fileId=id_folder,
            fields="id,name,mimeType"
        ).execute()
        if folder["mimeType"] != "application/vnd.google-apps.folder":
            messages.error(request, "El ID no corresponde a una carpeta.")
            return redirect("list_drive_files")
        # Buscar el primer archivo docx en la carpeta y se lo damos a register
        docx_files = service.files().list(
            q=f"'{folder['id']}' in parents and mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'",
            fields="files(id, name)",
            orderBy="name"
        ).execute().get("files", [])
        register = docx_files[0]["id"] if docx_files else None
        
        excel_files = service.files().list(
            q=f"'{folder['id']}' in parents and (mimeType='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet' or mimeType='application/vnd.ms-excel')",
            fields="files(id, name)",
            orderBy="name"
        ).execute().get("files", [])
        
        excel = excel_files[0]["id"] if excel_files else None
        
        drive_folder = Folder.objects.update_or_create(
            id=1,
            defaults={
                "id_folder": folder["id"],
                "name": folder["name"],
                "register_id": register if register else '0',
                "ex_id": excel if excel else '0'
                }
        )
        messages.success(request, f"Carpeta '{folder['name']}' seleccionada como activa.")
        return redirect("index")

    else:
        return redirect("drive_auth_init")

SCOPES = ['https://www.googleapis.com/auth/drive']
DRIVE_CLIENT = os.environ.get("DRIVE_CLIENT")
DRIVE_SECRET = os.environ.get("DRIVE_SECRET")
DRIVE_REDIRECT_URI = os.environ.get("DRIVE_REDIRECT_URI")

def drive_auth_init(request):
    client = DRIVE_CLIENT
    secret = DRIVE_SECRET
    redirect_uri = DRIVE_REDIRECT_URI
    flow = Flow.from_client_config(
        {
            "web": {
                "client_id": client,
                "client_secret": secret,
                "auth_uri": "https://accounts.google.com/o/oauth2/auth",
                "token_uri": "https://oauth2.googleapis.com/token",
                "redirect_uris": redirect_uri,
            }
        },
        scopes=SCOPES,
    )
    flow.redirect_uri = redirect_uri

    authorization_url, state = flow.authorization_url(
        access_type="offline",
        include_granted_scopes="true"
    )
    request.session['oauth_state'] = state
    return redirect(authorization_url)

from google.oauth2.credentials import Credentials

def drive_auth_callback(request):
    # Recupera el estado guardado al iniciar OAuth
    state = request.session.get('oauth_state')
    client = DRIVE_CLIENT
    secret = DRIVE_SECRET
    redirect_uri = DRIVE_REDIRECT_URI    
    flow = Flow.from_client_config(
        {
            "web": {
                "client_id": client,
                "client_secret": secret,
                "auth_uri": "https://accounts.google.com/o/oauth2/auth",
                "token_uri": "https://oauth2.googleapis.com/token",
                "redirect_uris": redirect_uri,
            }
        },
        scopes=SCOPES,
        state=state
    )
    flow.redirect_uri = os.environ["DRIVE_REDIRECT_URI"]

    # Construye la URL completa con parámetros que envió Google
    authorization_response = request.build_absolute_uri()
    flow.fetch_token(authorization_response=authorization_response)

    # Credenciales que permiten acceder a Google Drive
    credentials = flow.credentials

    # Guardarlas en la sesión de Django
    request.session['credentials'] = {
        'token': credentials.token,
        'refresh_token': credentials.refresh_token,
        'token_uri': credentials.token_uri,
        'client_id': credentials.client_id,
        'client_secret': credentials.client_secret,
        'scopes': credentials.scopes
    }

    # Redirigir a la vista que lista archivos o a otra página
    return redirect('index')

    # Construye la URL completa con parámetros que envió Google
    authorization_response = request.build_absolute_uri()
    flow.fetch_token(authorization_response=authorization_response)

    # Credenciales que permiten acceder a Google Drive
    credentials = flow.credentials

    # Guardarlas en la sesión de Django
    request.session['credentials'] = {
        'token': credentials.token,
        'refresh_token': credentials.refresh_token,
        'token_uri': credentials.token_uri,
        'client_id': credentials.client_id,
        'client_secret': credentials.client_secret,
        'scopes': credentials.scopes
    }

    # Redirigir a la vista que lista archivos o a otra página
    return redirect('index')

